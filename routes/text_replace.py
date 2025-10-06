from __future__ import annotations

import re
import os
import uuid
import zipfile
import threading
import time
from datetime import datetime, timedelta
from pathlib import Path
from tempfile import NamedTemporaryFile, TemporaryDirectory
from typing import List, Optional, Dict, Any

import pandas as pd
from fastapi import APIRouter, UploadFile, File, HTTPException, Query, Header, Form, Depends
from fastapi.responses import FileResponse
from openpyxl import load_workbook
from docx import Document
from sqlalchemy.orm import Session
from docx.oxml.ns import qn
from docx.text.run import Run

from ..config import ADMIN_TOKEN, TEXT_REPLACE_FILE_RETENTION_DAYS, TEXT_REPLACE_HISTORY_RETENTION_DAYS
from ..database import get_db, init_db
from ..models import TextReplaceHistory

router = APIRouter(prefix="/text-replace", tags=["Text Replace"])

_processed_zips: Dict[str, Dict[str, Any]] = {}
_zip_lock = threading.Lock()

try:
    init_db()
except Exception as e:
    print(f"Warning: Failed to initialize database: {e}")

def _cleanup_old_zips():
    """Clean up ZIP files based on keep_until time (thread-safe)"""
    current_time = time.time()
    to_remove = []
    
    for zip_id, info in _processed_zips.items():
        if current_time > info.get('keep_until', info['created_at'] + 3600):
            try:
                if os.path.exists(info['path']):
                    os.unlink(info['path'])
            except Exception as e:
                print(f"Warning: Failed to delete old ZIP {zip_id}: {e}")
            to_remove.append(zip_id)
    
    for zip_id in to_remove:
        _processed_zips.pop(zip_id, None)
    
    if to_remove:
        print(f"Cleaned up {len(to_remove)} old ZIP files")

def _require_admin(
    x_admin_token: Optional[str] = Header(None, convert_underscores=False),
    x_admin_token_alt: Optional[str] = Header(None, alias="X_Admin_Token"),
    authorization: Optional[str] = Header(None),
    token: Optional[str] = None,
):
    def _extract_admin_token(x_admin_token, x_admin_token_alt, authorization, token_qs):
        if x_admin_token and x_admin_token.strip():
            return x_admin_token.strip()
        if x_admin_token_alt and x_admin_token_alt.strip():
            return x_admin_token_alt.strip()
        if authorization and authorization.lower().startswith("bearer "):
            return authorization.split(" ", 1)[1].strip()
        if token_qs and token_qs.strip():
            return token_qs.strip()
        return ""
    
    if not ADMIN_TOKEN:
        raise HTTPException(status_code=500, detail="Admin token not configured")
    
    client_token = _extract_admin_token(x_admin_token, x_admin_token_alt, authorization, token)
    
    import hmac
    if not client_token or not hmac.compare_digest(client_token, ADMIN_TOKEN):
        raise HTTPException(status_code=401, detail="invalid admin token")

@router.post("/process")
async def text_replace(
    find_text: str = Form(...),
    replace_text: str = Form(...),
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    """Replace text in multiple Word/Excel files with partial success support"""
    print(f"Text Replace Request: find='{find_text}', replace='{replace_text}', files={len(files)}")
    
    if not find_text.strip() or not replace_text.strip():
        raise HTTPException(status_code=400, detail="Find and replace text cannot be empty")
    
    if len(files) > 50:
        raise HTTPException(status_code=400, detail="Maximum 50 files allowed")
    
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    import shutil as disk_util
    free_space = disk_util.disk_usage('/tmp').free
    if free_space < 100 * 1024 * 1024:
        raise HTTPException(status_code=507, detail="Insufficient disk space")
    
    results = {
        "successful": [],
        "failed": [],
        "total_files": len(files),
        "processed_count": 0
    }
    
    try:
        zip_temp_file = NamedTemporaryFile(delete=False, suffix=".zip")
        zip_path = zip_temp_file.name
        zip_temp_file.close()
        
        with TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            processed_files = []
            
            for i, file in enumerate(files):
                file_result = {
                    "original_name": file.filename or f"file_{i}",
                    "index": i,
                    "error": None
                }
                
                try:
                    if not file.filename:
                        raise ValueError("Empty filename")
                    
                    if file.size and file.size > MAX_FILE_SIZE:
                        raise ValueError(f"File exceeds 10MB limit ({file.size} bytes)")
                    
                    if file.size == 0:
                        raise ValueError("Empty file")
                    
                    ext = file.filename.lower().split('.')[-1]
                    if ext not in ['docx', 'xlsx']:
                        raise ValueError(f"Unsupported file type: .{ext}")
                    
                    import re
                    
                    clean_name = re.sub(r'[^a-zA-Z0-9._-]', '_', file.filename)
                    clean_name = clean_name[:50]
                    
                    if not clean_name or clean_name.startswith('.') or '..' in clean_name:
                        clean_name = f"file_{i:03d}.{ext}"
                    
                    if not clean_name.endswith(f'.{ext}'):
                        clean_name = f"{clean_name.split('.')[0]}.{ext}"
                    
                    original_path = temp_path / clean_name
                    
                    counter = 1
                    while original_path.exists():
                        name_part = clean_name.rsplit('.', 1)[0]
                        original_path = temp_path / f"{name_part}_{counter}.{ext}"
                        counter += 1
                    
                    content = await file.read(MAX_FILE_SIZE + 1)
                    if len(content) > MAX_FILE_SIZE:
                        raise ValueError(f"File exceeds size limit: {len(content)} bytes")
                    if not content:
                        raise ValueError("File content is empty")
                    
                    if ext == 'docx' and not content.startswith(b'PK'):
                        raise ValueError("Invalid DOCX file format")
                    elif ext == 'xlsx' and not content.startswith(b'PK'):
                        raise ValueError("Invalid XLSX file format")
                    
                    with open(original_path, 'wb') as f:
                        f.write(content)
                    
                    print(f"Processing file {i+1}/{len(files)}: {file.filename} ({ext})")
                    
                    if ext == 'docx':
                        replacements = _replace_text_in_docx_safe(original_path, find_text, replace_text)
                    elif ext == 'xlsx':
                        replacements = _replace_text_in_xlsx_safe(original_path, find_text, replace_text)
                    else:
                        raise ValueError(f"Unsupported extension: {ext}")
                    
                    processed_files.append((original_path, file.filename))
                    file_result["replacements"] = replacements
                    file_result["status"] = "success" if replacements > 0 else "no_matches"
                    file_result["message"] = f"{replacements} replacements made" if replacements > 0 else f"No matches found for '{find_text}'"
                    results["successful"].append(file_result)
                    results["processed_count"] += 1
                    
                    print(f"Successfully processed: {file.filename} ({replacements} replacements)")
                    
                except Exception as e:
                    error_msg = str(e)
                    print(f"Error processing {file.filename}: {error_msg}")
                    file_result["error"] = error_msg
                    results["failed"].append(file_result)
            
            if not processed_files:
                raise HTTPException(status_code=400, detail="No files were successfully processed")
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for temp_file_path, original_filename in processed_files:
                    zipf.write(temp_file_path, original_filename)
        
        zip_id = str(uuid.uuid4())
        current_time = time.time()
        
        with _zip_lock:
            _processed_zips[zip_id] = {
                'path': zip_path,
                'created_at': current_time,
                'keep_until': current_time + (TEXT_REPLACE_FILE_RETENTION_DAYS * 86400)
            }
            _cleanup_old_zips()
        
        try:
            from ..database import SessionLocal
            db = SessionLocal()
            try:
                total_replacements = sum(f.get("replacements", 0) for f in results["successful"])
                no_matches = len([f for f in results["successful"] if f.get("replacements", 0) == 0])
                with_matches = len([f for f in results["successful"] if f.get("replacements", 0) > 0])
                
                history_record = TextReplaceHistory(
                    zip_id=zip_id,
                    find_text=find_text,
                    replace_text=replace_text,
                    total_files=results["total_files"],
                    successful=len(results["successful"]),
                    failed=len(results["failed"]),
                    success_rate=f"{len(results['successful'])/len(files)*100:.1f}%",
                    total_replacements=total_replacements,
                    files_with_matches=with_matches,
                    files_no_matches=no_matches,
                    zip_path=zip_path,
                    expires_at=datetime.now() + timedelta(days=TEXT_REPLACE_FILE_RETENTION_DAYS)
                )
                db.add(history_record)
                db.commit()
            finally:
                db.close()
        except Exception as e:
            print(f"Warning: Failed to save history to database: {e}")
        
        print(f"Text replacement completed. Processed {len(processed_files)}/{len(files)} files.")
        
        return {
            "zip_id": zip_id,
            "results": results,
            "summary": {
                "total_files": results["total_files"],
                "successful": len(results["successful"]),
                "failed": len(results["failed"]),
                "success_rate": f"{len(results['successful'])/len(files)*100:.1f}%"
            }
        }
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Text replacement failed: {str(e)}")

@router.get("/download/{zip_id}")
def download_zip_file(zip_id: str):
    """Download a processed ZIP file by ID with thread-safe cleanup"""
    with _zip_lock:
        if zip_id not in _processed_zips:
            raise HTTPException(status_code=404, detail="ZIP file not found")
        
        zip_info = _processed_zips[zip_id]
        zip_path = zip_info['path']
        
        if not os.path.exists(zip_path):
            del _processed_zips[zip_id]
            raise HTTPException(status_code=404, detail="ZIP file no longer exists")
    
    response = FileResponse(
        path=zip_path,
        media_type="application/zip",
        filename=f"replaced_files_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.zip"
    )
    
    return response

@router.get("/history")
def get_history(db: Session = Depends(get_db)):
    """Get text replacement history from database"""
    try:
        records = db.query(TextReplaceHistory).order_by(
            TextReplaceHistory.created_at.desc()
        ).limit(50).all()
        
        history_list = []
        current_time = datetime.now()
        
        for record in records:
            zip_available = (
                record.expires_at > current_time and 
                os.path.exists(record.zip_path)
            )
            
            if record.zip_available and not zip_available:
                record.zip_available = False
                db.commit()
            

            
            history_list.append({
                'id': record.id,
                'zip_id': record.zip_id,
                'find_text': record.find_text,
                'replace_text': record.replace_text,
                'total_files': record.total_files,
                'successful': record.successful,
                'failed': record.failed,
                'success_rate': record.success_rate,
                'total_replacements': getattr(record, 'total_replacements', 0),
                'files_with_matches': getattr(record, 'files_with_matches', 0),
                'files_no_matches': getattr(record, 'files_no_matches', 0),
                'timestamp': (record.created_at + timedelta(hours=7)).strftime('%Y-%m-%d %H:%M:%S'),
                'zip_available': zip_available
            })
        
        return {"history": history_list}
        
    except Exception as e:
        print(f"Error loading history: {e}")
        return {"history": []}

@router.get("/history/{zip_id}/download")
def download_history_zip(zip_id: str, db: Session = Depends(get_db)):
    """Download ZIP from history"""
    try:
        record = db.query(TextReplaceHistory).filter(
            TextReplaceHistory.zip_id == zip_id
        ).first()
        
        if not record:
            raise HTTPException(status_code=404, detail="ZIP file not found in history")
        
        if datetime.now() > record.expires_at:
            record.zip_available = False
            db.commit()
            raise HTTPException(status_code=410, detail="ZIP file has expired")
        
        if not os.path.exists(record.zip_path):
            record.zip_available = False
            db.commit()
            raise HTTPException(status_code=404, detail="ZIP file no longer exists")
        
        return FileResponse(
            path=record.zip_path,
            media_type="application/zip",
            filename=f"history_files_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.zip"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error downloading history ZIP: {e}")
        raise HTTPException(status_code=500, detail="Failed to download file")

@router.post("/admin/cleanup")
def cleanup_all_zips(
    x_admin_token: Optional[str] = Header(None, convert_underscores=False),
    x_admin_token_alt: Optional[str] = Header(None, alias="X_Admin_Token"),
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Clean up expired ZIP files (admin only)"""
    _require_admin(x_admin_token, x_admin_token_alt, authorization, None)
    
    try:
        current_time = datetime.now()
        
        expired_records = db.query(TextReplaceHistory).filter(
            TextReplaceHistory.expires_at < current_time
        ).all()
        
        old_cutoff = current_time - timedelta(days=TEXT_REPLACE_HISTORY_RETENTION_DAYS)
        old_records = db.query(TextReplaceHistory).filter(
            TextReplaceHistory.created_at < old_cutoff
        ).all()
        
        cleaned_count = 0
        deleted_records = 0
        
        for record in expired_records:
            try:
                if os.path.exists(record.zip_path):
                    os.unlink(record.zip_path)
                record.zip_available = False
                cleaned_count += 1
            except Exception as e:
                print(f"Warning: Failed to delete ZIP {record.zip_id}: {e}")
        
        for record in old_records:
            try:
                if os.path.exists(record.zip_path):
                    os.unlink(record.zip_path)
                db.delete(record)
                deleted_records += 1
            except Exception as e:
                print(f"Warning: Failed to delete old record {record.zip_id}: {e}")
        
        db.commit()
        return {
            "ok": True, 
            "cleaned_files": cleaned_count,
            "deleted_records": deleted_records,
            "file_retention_days": TEXT_REPLACE_FILE_RETENTION_DAYS,
            "history_retention_days": TEXT_REPLACE_HISTORY_RETENTION_DAYS
        }
        
    except Exception as e:
        print(f"Error during cleanup: {e}")
        return {"ok": False, "error": str(e)}

@router.get("/admin/storage-status")
def get_storage_status(
    x_admin_token: Optional[str] = Header(None, convert_underscores=False),
    x_admin_token_alt: Optional[str] = Header(None, alias="X_Admin_Token"),
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Get storage status and retention info (admin only)"""
    _require_admin(x_admin_token, x_admin_token_alt, authorization, None)
    
    try:
        current_time = datetime.now()
        
        total_records = db.query(TextReplaceHistory).count()
        available_files = db.query(TextReplaceHistory).filter(
            TextReplaceHistory.zip_available == True,
            TextReplaceHistory.expires_at > current_time
        ).count()
        expired_files = db.query(TextReplaceHistory).filter(
            TextReplaceHistory.expires_at < current_time
        ).count()
        
        import shutil
        disk_usage = shutil.disk_usage('/tmp')
        
        return {
            "total_records": total_records,
            "available_files": available_files,
            "expired_files": expired_files,
            "file_retention_days": TEXT_REPLACE_FILE_RETENTION_DAYS,
            "history_retention_days": TEXT_REPLACE_HISTORY_RETENTION_DAYS,
            "disk_usage": {
                "total_gb": round(disk_usage.total / (1024**3), 2),
                "used_gb": round((disk_usage.total - disk_usage.free) / (1024**3), 2),
                "free_gb": round(disk_usage.free / (1024**3), 2),
                "usage_percent": round((disk_usage.total - disk_usage.free) / disk_usage.total * 100, 1)
            }
        }
        
    except Exception as e:
        return {"error": str(e)}

@router.post("/admin/delete-selected")
def delete_selected_records(
    record_ids: List[int] = Form(...),
    x_admin_token: Optional[str] = Header(None, convert_underscores=False),
    x_admin_token_alt: Optional[str] = Header(None, alias="X_Admin_Token"),
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
):
    """Delete selected history records (admin only)"""
    _require_admin(x_admin_token, x_admin_token_alt, authorization, None)
    
    try:
        deleted_count = 0
        for record_id in record_ids:
            record = db.query(TextReplaceHistory).filter(
                TextReplaceHistory.id == record_id
            ).first()
            
            if record:
                try:
                    if os.path.exists(record.zip_path):
                        os.unlink(record.zip_path)
                    db.delete(record)
                    deleted_count += 1
                except Exception as e:
                    print(f"Warning: Failed to delete record {record_id}: {e}")
        
        db.commit()
        return {"ok": True, "deleted_count": deleted_count}
        
    except Exception as e:
        return {"ok": False, "error": str(e)}

def _replace_text_in_docx_safe(file_path: Path, find_text: str, replace_text: str) -> int:
    """Replace text in DOCX file with cross-run replacement support and enforce TH SarabunPSK font"""
    print(f"Processing DOCX: {file_path.name}")
    print(f"Looking for: '{find_text}' -> '{replace_text}'")

    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f"Cannot open DOCX file: {e}")



    replacements_made = 0

    def _clone_run_format(src: Run, dst: Run):
        """Copy all formatting from src run to dst run"""
        try:
            dst.style = src.style
            dst.bold = src.bold
            dst.italic = src.italic
            dst.underline = src.underline
            
            # Preserve original font
            if src.font.name:
                dst.font.name = src.font.name
            
            # Copy font properties
            if src.font.size:
                dst.font.size = src.font.size
            if src.font.color and src.font.color.rgb:
                dst.font.color.rgb = src.font.color.rgb
            
            dst.font.highlight_color = src.font.highlight_color
            dst.font.all_caps = src.font.all_caps
            dst.font.small_caps = src.font.small_caps
            
            # Copy font family settings
            src_rPr = src._element.rPr
            dst_rPr = dst._element.rPr
            if src_rPr is not None and dst_rPr is not None:
                src_fonts = src_rPr.rFonts
                if src_fonts is not None:
                    dst_rPr.rFonts.set(qn('w:ascii'), src_fonts.get(qn('w:ascii')))
                    dst_rPr.rFonts.set(qn('w:hAnsi'), src_fonts.get(qn('w:hAnsi')))
                    dst_rPr.rFonts.set(qn('w:eastAsia'), src_fonts.get(qn('w:eastAsia')))
        except Exception as e:
            print(f"Warning: Could not copy all formatting: {e}")

    def _gather_paragraph_text_and_run_info(paragraph):
        """Return concatenated paragraph text and a list of run info objects"""
        full_text = ""
        run_info = []
        for r in paragraph.runs:
            t = r.text or ""
            full_text += t
            run_info.append({'text': t, 'run': r})
        return full_text, run_info

    def _replace_in_paragraph_preserve_format(paragraph, needle, repl):
        nonlocal replacements_made
        if not paragraph.runs or not needle:
            return 0
        
        full_text, run_info = _gather_paragraph_text_and_run_info(paragraph)
        if needle not in full_text:
            return 0

        count = full_text.count(needle)
        new_text = full_text.replace(needle, repl)
        
        # Get first run's formatting as template
        template_run = paragraph.runs[0] if paragraph.runs else None
        
        # Clear all runs
        for run in paragraph.runs[:]:
            paragraph._element.remove(run._element)
        
        # Add new run with preserved formatting
        new_run = paragraph.add_run(new_text)
        if template_run:
            _clone_run_format(template_run, new_run)
        
        replacements_made += count
        return count

    # Replace in body
    for p in doc.paragraphs:
        _replace_in_paragraph_preserve_format(p, find_text, replace_text)

    # Replace in tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph_preserve_format(p, find_text, replace_text)

    # Replace in headers/footers
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                _replace_in_paragraph_preserve_format(p, find_text, replace_text)
        if section.footer:
            for p in section.footer.paragraphs:
                _replace_in_paragraph_preserve_format(p, find_text, replace_text)

    doc.save(file_path)
    print(f"*** FINAL RESULT: {replacements_made} total replacements made in {file_path.name} ***")
    return replacements_made

def _replace_text_in_xlsx_safe(file_path: Path, find_text: str, replace_text: str) -> int:
    """Replace text in XLSX file with improved data type handling"""
    print(f"Processing XLSX: {file_path.name}")
    
    try:
        wb = load_workbook(file_path)
    except Exception as e:
        raise ValueError(f"Cannot open XLSX file: {e}")
    
    replacements_made = 0
    
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            for row in ws.iter_rows():
                for cell in row:
                    try:
                        if cell.value is None:
                            continue
                        
                        if hasattr(cell, 'data_type') and cell.data_type == 'f':
                            continue
                        
                        cell_str = str(cell.value)
                        
                        if find_text in cell_str:
                            old_value = cell.value
                            old_type = type(cell.value)
                            
                            new_str = cell_str.replace(find_text, replace_text)
                            
                            try:
                                if isinstance(old_value, str):
                                    cell.value = new_str
                                elif isinstance(old_value, int):
                                    try:
                                        cell.value = int(new_str)
                                    except ValueError:
                                        try:
                                            cell.value = float(new_str)
                                        except ValueError:
                                            cell.value = new_str
                                elif isinstance(old_value, float):
                                    try:
                                        cell.value = float(new_str)
                                    except ValueError:
                                        cell.value = new_str
                                else:
                                    cell.value = new_str
                                
                                if old_value != cell.value:
                                    replacements_made += 1
                                    print(f"  Replaced '{old_value}' -> '{cell.value}' in {sheet_name}")
                                    
                            except Exception as e:
                                print(f"  Warning: Error updating cell {cell.coordinate}: {e}")
                                
                    except Exception as e:
                        print(f"  Warning: Error processing cell: {e}")
                        continue
        
        wb.save(file_path)
        print(f"*** FINAL RESULT: {replacements_made} total replacements made in {file_path.name} ***")
        return replacements_made
        
    except Exception as e:
        raise ValueError(f"Error processing XLSX: {e}")